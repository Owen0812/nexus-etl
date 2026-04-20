"""Word (.docx) and HTML extractors — produce the same output format as vision_extractor."""
from __future__ import annotations

from pathlib import Path

from backend.agents.state import DocumentState


def word_extractor_node(state: DocumentState) -> dict:
    """Extract text and tables from a .docx file using python-docx."""
    from docx import Document as DocxDocument  # lazy import — optional dependency
    from docx.oxml.ns import qn

    doc = DocxDocument(state["file_path"])
    pages: list[dict] = []
    tables: list[dict] = []

    # Word has no real pages; treat each paragraph section as a virtual page
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    pages.append({"page_num": 0, "text": full_text, "width": None, "height": None})

    for tbl_idx, tbl in enumerate(doc.tables):
        rows = []
        for row in tbl.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({
            "page_num": 0,
            "table_index": tbl_idx,
            "data": rows,
            "vision_description": None,
        })

    return {
        "raw_pages": pages,
        "extracted_tables": tables,
        "extracted_images": [],
        "stages_completed": ["vision_extractor"],  # unified stage key
    }


def html_extractor_node(state: DocumentState) -> dict:
    """Extract text and tables from an HTML file using BeautifulSoup."""
    from bs4 import BeautifulSoup

    html = Path(state["file_path"]).read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html, "lxml")

    # Remove script/style noise
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    full_text = soup.get_text(separator="\n", strip=True)
    pages = [{"page_num": 0, "text": full_text, "width": None, "height": None}]

    tables: list[dict] = []
    for tbl_idx, tbl in enumerate(soup.find_all("table")):
        rows = []
        for tr in tbl.find_all("tr"):
            cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
            if cells:
                rows.append(cells)
        if rows:
            tables.append({
                "page_num": 0,
                "table_index": tbl_idx,
                "data": rows,
                "vision_description": None,
            })

    return {
        "raw_pages": pages,
        "extracted_tables": tables,
        "extracted_images": [],
        "stages_completed": ["vision_extractor"],
    }
